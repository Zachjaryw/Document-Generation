from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.style import WD_STYLE_TYPE

class Elements:
    
    def text(document:object,text:str,fontsize:int,color:list = [0,0,0],bold = False,italics = False,underline = False):
        '''
        This method creates text on the document
        @param document (Document) - python-docx document object
        @param text (str) - the text to be added to the page
        @fontsize (int) - the size of the text in pts
        @color (list) - the color of the text on the page in RGB list (default - [0,0,0] black)
        @return documetn (Document) - python-docx document object updated
        '''
        paragraph = document.add_paragraph().add_run(text)
        font = paragraph.font
        font.name = "Helvetica"
        font.size = Pt(fontsize)
        font.color.rgb = RGBColor(color[0],color[1],color[2])
        font.bold = bold
        font.italic = italics
        font.underline = underline
        return document

    
    def table(document:object,text:list,fontsize:int,color:list = [0,0,0],bold:any = False):
        '''
        This method creates a table on the document with the given size and criteria
        @param document (Document) - python-docx document object
        @param text (2d list) - a 2d list containing the text that will be added to the table
            formatted with each row as its own list within the list of lists
        @param fontsize (int) - the font size the items in the text box will be
        @color (list) - the color of the text on the page in RGB list (default - [0,0,0] black)
        @return document (Document) - python-docx document object updated
        '''
        rows,cols = len(text),len(text[0])
        table = document.add_table(rows,cols)
        for row in range(len(text)):
            for column in range(len(text[row])):
                paragraph = table.cell(row,column).add_paragraph().add_run(text = str(text[row][column]))
                font = paragraph.font
                font.name = "Helvetica"
                font.size = Pt(fontsize)
                font.color.rgb = RGBColor(color[0],color[1],color[2])
                if type(bold) == list:
                    font.bold = bold[row][column]
                else:
                    if row == 0:
                        font.bold = True
                    else:
                        font.bold = False
        return document
    
    def image(document:object,filepath:str):
        '''
        This method adds an image to the document with the given filepath
        @param document (Document) - python-docx document object
        @param filepath (str) - string for filepath of image
        @return document (Document) - updated document
        '''
        document.add_picture(filepath)
        return document